import os
import json
from datetime import datetime
import shutil
import sys
import configparser

import PyPDF2
from docx import Document
from docx.shared import Inches
import pytesseract
from PIL import Image
import pypandoc
from pdf2image import convert_from_path

OUTPUT_SNAPSHOTS_DOCX_TEMPLATE = "Extracted Texts with Snapshots_{}.docx"
OUTPUT_COMBINED_TEXT_DOCX_TEMPLATE = "Extracted Texts_{}.docx"
OUTPUT_TEXT_FOLDER_TEMPLATE = "Extracted Texts for Iteration_{}"

def extract_text_from_pdf(pdf_path):
    text = ""
    try:
        with open(pdf_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                text += page.extract_text() or ""
    except Exception as e:
        print(f"Error extracting text from PDF {pdf_path}: {e}")
    return text

def extract_text_from_image(image_path):
    text = ""
    try:
        text = pytesseract.image_to_string(Image.open(image_path))
    except pytesseract.TesseractNotFoundError:
        print("Error: Tesseract not found. Please install it and add to PATH to process images.")
    except Exception as e:
        print(f"Error extracting text from image {image_path}: {e}")
    return text

def extract_text_from_docx(docx_path):
    text = ""
    try:
        doc = Document(docx_path)
        for para in doc.paragraphs:
            text += para.text + "\n"
    except Exception as e:
        print(f"Error extracting text from DOCX {docx_path}: {e}")
    return text

def extract_text_from_doc(doc_path):
    try:
        return pypandoc.convert_file(doc_path, 'plain')
    except Exception as e:
        print(f"Error converting .doc file {doc_path} with pypandoc: {e}")
        return ""

def get_file_text(file_path):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        return extract_text_from_pdf(file_path)
    elif ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
        return extract_text_from_image(file_path)
    elif ext == '.docx':
        return extract_text_from_docx(file_path)
    elif ext == '.doc':
        return extract_text_from_doc(file_path)
    else:
        print(f"Skipping unsupported file type for text extraction: {os.path.basename(file_path)}")
        return ""

def create_pdf_snapshot(pdf_path, output_dir):
    try:
        images = convert_from_path(pdf_path, first_page=1, last_page=1, dpi=150)
        if images:
            image_path = os.path.join(output_dir, f"{os.path.basename(pdf_path)}.png")
            images[0].save(image_path, 'PNG')
            return image_path
    except Exception as e:
        print(f"Error generating PDF snapshot for {os.path.basename(pdf_path)}: {e}. (Is Poppler installed?)")
    return None

def create_word_snapshot(word_path, output_dir):
    temp_pdf_path = os.path.join(output_dir, f"{os.path.basename(word_path)}.pdf")
    snapshot_path = None
    try:
        pypandoc.convert_file(word_path, 'pdf', outputfile=temp_pdf_path)
        
        if os.path.exists(temp_pdf_path):
            snapshot_path = create_pdf_snapshot(temp_pdf_path, output_dir)
            
    except Exception as e:
        print(f"Info: Could not generate snapshot for Word file '{os.path.basename(word_path)}'.")
        print("      This usually means a PDF engine like 'pdflatex' (from MacTeX/MiKTeX) is not installed.")
        print(f"      Pandoc error: {e}")
    finally:
        if os.path.exists(temp_pdf_path):
            os.remove(temp_pdf_path)
            
    return snapshot_path

def get_file_snapshot(file_path, temp_dir):
    ext = os.path.splitext(file_path)[1].lower()
    if ext == '.pdf':
        return create_pdf_snapshot(file_path, temp_dir)
    elif ext in ['.docx', '.doc']:
        return create_word_snapshot(file_path, temp_dir)
    elif ext in ['.jpg', '.jpeg', '.png', '.tiff', '.bmp']:
        return file_path # The file is its own snapshot
    else:
        return None

def create_snapshots_docx(extracted_data, output_path, root_folder):
    document = Document()
    document.add_heading('Raw Text Extraction with Visual Snapshots', level=1)
    document.add_paragraph(f"Report generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    temp_dir = os.path.join(root_folder, 'temp_snapshots_deleteme')
    os.makedirs(temp_dir, exist_ok=True)

    for entry in extracted_data:
        filename = entry.get('filename', 'N/A')
        text_content = entry.get('text', 'No text could be extracted.')
        file_path = entry.get('file_path', '')

        document.add_heading(f"File: {filename}", level=2)
        document.add_paragraph("--- Extracted Text ---").add_run().bold = True
        document.add_paragraph(text_content or "[No text extracted]")
        
        document.add_paragraph("\n--- Snapshot ---").add_run().bold = True
        snapshot_path = get_file_snapshot(file_path, temp_dir)
        if snapshot_path and os.path.exists(snapshot_path):
            try:
                document.add_picture(snapshot_path, width=Inches(6.0))
            except Exception as e:
                document.add_paragraph(f"[Error embedding snapshot: {e}]")
        else:
            document.add_paragraph("[Snapshot not available or could not be generated.]")
        document.add_page_break()

    try:
        document.save(output_path)
        print(f"\nSuccessfully saved snapshot report to: {output_path}")
    except Exception as e:
        print(f"\nError saving snapshot DOCX report: {e}")
    finally:
        shutil.rmtree(temp_dir) # Clean up temp snapshot folder

def create_combined_text_docx(extracted_data, output_path):
    document = Document()
    document.add_heading('Combined Raw Text Extraction', level=1)
    document.add_paragraph(f"Report generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    for entry in extracted_data:
        filename = entry.get('filename', 'N/A')
        text_content = entry.get('text', 'No text could be extracted.')
        document.add_heading(f"--- File: {filename} ---", level=2)
        document.add_paragraph(text_content or "[No text extracted]")
        document.add_paragraph("\n" + "="*80 + "\n")

    try:
        document.save(output_path)
        print(f"Successfully saved combined text report to: {output_path}")
    except Exception as e:
        print(f"Error saving combined text DOCX report: {e}")

def create_individual_text_files(extracted_data, output_dir, combine_count):
    if os.path.exists(output_dir):
        shutil.rmtree(output_dir)
    os.makedirs(output_dir, exist_ok=True)
    
    print(f"Creating individual/combined text files in: {output_dir}")

    if combine_count <= 0:
        print("Warning: TEXT_FILE_COMBINATION_COUNT must be 1 or greater. Defaulting to 1.")
        combine_count = 1

    num_entries = len(extracted_data)
    for i in range(0, num_entries, combine_count):
        chunk = extracted_data[i:i + combine_count]
        
        # Create a filename for the combined text file.
        base_filenames = [os.path.splitext(entry['filename'])[0] for entry in chunk]
        out_filename = "_and_".join(base_filenames) + ".txt"
        out_filepath = os.path.join(output_dir, out_filename)

        try:
            with open(out_filepath, 'w', encoding='utf-8') as f:
                for entry in chunk:
                    filename = entry.get('filename', 'N/A')
                    text_content = entry.get('text', '[No text extracted]')
                    f.write(f"--- File: {filename} ---\n\n")
                    f.write(text_content)
                    f.write("\n\n" + "="*80 + "\n\n")
        except Exception as e:
            print(f"Error saving text file '{out_filename}': {e}")

def main():
    config = configparser.ConfigParser()
    config.read('config.ini')
    text_file_combination_count = config.getint('Settings', 'combination_count', fallback=1)

    target_folder_path = input("Please enter the path to the folder to process: ").strip()
    
    if not os.path.isdir(target_folder_path):
        print(f"Error: The provided path '{target_folder_path}' is not a valid directory.")
        sys.exit(1)

    output_directory = os.path.dirname(target_folder_path)
    target_folder_name = os.path.basename(target_folder_path)

    metadata_path = os.path.join(target_folder_path, 'files_metadata.json')
    if not os.path.exists(metadata_path):
        print(f"Error: 'files_metadata.json' not found in '{target_folder_path}'.")
        sys.exit(1)

    with open(metadata_path, 'r', encoding='utf-8') as f:
        metadata = json.load(f)

    all_extracted_data = []
    print(f"\nProcessing files from: {target_folder_path}")

    for file_id, file_info in metadata.items():
        filename = file_info.get('filename')
        file_path = file_info.get('local_path')
        
        if not (filename and file_path and os.path.isfile(file_path)):
            print(f"Warning: Skipping invalid entry (ID: {file_id}). File may be missing.")
            continue

        print(f"  - Processing: {filename}")
        document_text = get_file_text(file_path)
        
        all_extracted_data.append({
            "filename": filename,
            "text": document_text,
            "file_path": file_path,
        })

    print("\n--- All Files Processed ---")

    if not all_extracted_data:
        print("No text was extracted from any files. No output files will be generated.")
        return

    snapshot_report_name = OUTPUT_SNAPSHOTS_DOCX_TEMPLATE.format(target_folder_name)
    snapshot_report_path = os.path.join(output_directory, snapshot_report_name)
    create_snapshots_docx(all_extracted_data, snapshot_report_path, output_directory)

    combined_text_name = OUTPUT_COMBINED_TEXT_DOCX_TEMPLATE.format(target_folder_name)
    combined_text_path = os.path.join(output_directory, combined_text_name)
    create_combined_text_docx(all_extracted_data, combined_text_path)

    text_folder_name = OUTPUT_TEXT_FOLDER_TEMPLATE.format(target_folder_name)
    text_folder_path = os.path.join(output_directory, text_folder_name)
    create_individual_text_files(all_extracted_data, text_folder_path, text_file_combination_count)

if __name__ == '__main__':
    main()